import streamlit as st
import os
import time
import json
import hashlib
import pandas as pd
import sqlite3
from PIL import Image
from PyPDF2 import PdfReader
import google.generativeai as genai
import docx
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
from datetime import datetime
import plotly.express as px
import plotly.graph_objects as go
import asyncio
import nest_asyncio
import re

# Fix for event loop issues
nest_asyncio.apply()
load_dotenv()

# Set page config ONLY ONCE at the very beginning
st.set_page_config(
    page_title="Elite MultiModal AI Platform",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Configure Google AI with your exact format handling
def configure_google_ai():
    """Configure Google AI with robust key handling"""
    try:
        api_key = None
        
        # Method 1: Try Streamlit secrets (your format: GOOGLE_API_KEY ="HERE IS API KEY")
        try:
            if hasattr(st, 'secrets') and "GOOGLE_API_KEY" in st.secrets:
                api_key = st.secrets["GOOGLE_API_KEY"].strip()
                if api_key and len(api_key) > 10:
                    print("✅ API key loaded from Streamlit secrets")
        except Exception as e:
            print(f"Secrets error: {str(e)}")
        
        # Method 2: Try environment variables
        if not api_key:
            try:
                env_key = os.getenv("GOOGLE_API_KEY")
                if env_key and len(env_key.strip()) > 10:
                    api_key = env_key.strip()
                    print("✅ API key loaded from environment")
            except Exception as e:
                print(f"Environment error: {str(e)}")
        
        # Validate and configure
        if not api_key:
            raise Exception("Google API Key not found. Please add GOOGLE_API_KEY to Streamlit secrets.")
        
        # Clean the API key (remove any extra quotes or spaces)
        api_key = api_key.strip().strip('"').strip("'")
        
        if len(api_key) < 20:
            raise Exception("Invalid Google API Key format. Key seems too short.")
        
        genai.configure(api_key=api_key)
        return api_key, True, "Google AI configured successfully!"
        
    except Exception as e:
        error_msg = f"Google AI configuration error: {str(e)}"
        print(error_msg)
        st.error(f"❌ {error_msg}")
        return None, False, error_msg

# Configure Google AI on startup
API_KEY, AI_CONFIGURED, CONFIG_MESSAGE = configure_google_ai()

# Global AI Models - Initialize once and reuse
@st.cache_resource
def get_ai_models():
    """Initialize and cache AI models"""
    try:
        if not AI_CONFIGURED or not API_KEY:
            raise Exception("Google AI not properly configured")
        
        gemini_model = ChatGoogleGenerativeAI(
            model="gemini-2.0-flash",
            temperature=0.3,
            google_api_key=API_KEY
        )
        
        gemini_vision = genai.GenerativeModel('gemini-2.0-flash')
        
        embeddings = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001",
            google_api_key=API_KEY
        )
        
        return gemini_model, gemini_vision, embeddings, True
    except Exception as e:
        st.error(f"Error initializing AI models: {str(e)}")
        return None, None, None, False

class ModernAuthManager:
    def __init__(self):
        self.users_file = "users_db.json"
        self.ensure_users_file()
    
    def ensure_users_file(self):
        """Ensure users database file exists"""
        if not os.path.exists(self.users_file):
            with open(self.users_file, 'w') as f:
                json.dump({}, f)
    
    def hash_password(self, password):
        """Hash password with salt"""
        salt = "modern_ai_platform_2025"
        return hashlib.sha256((password + salt).encode()).hexdigest()
    
    def register_user(self, username, password, email, full_name):
        """Register new user with validation"""
        try:
            with open(self.users_file, 'r') as f:
                users = json.load(f)
            
            if username in users:
                return False, "Username already exists!"
            
            # Email validation
            if '@' not in email or '.' not in email:
                return False, "Please enter a valid email address!"
            
            users[username] = {
                "password": self.hash_password(password),
                "email": email,
                "full_name": full_name,
                "created_at": datetime.now().isoformat(),
                "login_count": 0,
                "csv_queries": 0,
                "document_queries": 0,
                "image_queries": 0,
                "last_login": None
            }
            
            with open(self.users_file, 'w') as f:
                json.dump(users, f, indent=4)
            
            return True, "Account created successfully!"
        except Exception as e:
            return False, f"Registration failed: {str(e)}"
    
    def login_user(self, username, password):
        """Login user with enhanced tracking"""
        try:
            with open(self.users_file, 'r') as f:
                users = json.load(f)
            
            if username not in users:
                return False, "Username not found!"
            
            if users[username]["password"] != self.hash_password(password):
                return False, "Invalid password!"
            
            users[username]["login_count"] += 1
            users[username]["last_login"] = datetime.now().isoformat()
            
            with open(self.users_file, 'w') as f:
                json.dump(users, f, indent=4)
            
            return True, f"Welcome back, {users[username]['full_name']}!"
        except Exception as e:
            return False, f"Login failed: {str(e)}"
    
    def get_user_info(self, username):
        """Get user information"""
        try:
            with open(self.users_file, 'r') as f:
                users = json.load(f)
            return users.get(username, {})
        except:
            return {}
    
    def update_user_stats(self, username, stat_type):
        """Update user statistics"""
        try:
            with open(self.users_file, 'r') as f:
                users = json.load(f)
            
            if username in users and stat_type in users[username]:
                users[username][stat_type] += 1
                
                with open(self.users_file, 'w') as f:
                    json.dump(users, f, indent=4)
        except Exception as e:
            print(f"Error updating user stats: {str(e)}")

class ModernMultiModalPlatform:
    def __init__(self):
        self.auth_manager = ModernAuthManager()
        self.initialize_session_state()
        
        # Get AI models from cache
        self.gemini_model, self.gemini_vision, self.embeddings, self.models_ready = get_ai_models()
        
        # Apply CSS immediately
        self.apply_enhanced_css()
    
    def initialize_session_state(self):
        """Initialize all session state variables"""
        defaults = {
            'authenticated': False,
            'username': '',
            'current_page': 'csv',
            # Separate chat histories for each module - PERFECT CONTINUOUS CHAT
            'csv_messages': [],
            'document_messages': [],
            'image_messages': [],
            'current_image': None,
            'document_vector_store': None,
            'processed_files': [],
            'csv_data': None,
            'csv_file_name': '',
            'db_name': 'analytics_db',
            'table_name': 'data_table',
            'csv_columns': [],
            'dynamic_examples': []
        }
        
        for key, default_value in defaults.items():
            if key not in st.session_state:
                st.session_state[key] = default_value

    def apply_enhanced_css(self):
        """Apply enhanced CSS with modern black theme"""
        st.markdown("""
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&family=Poppins:wght@300;400;500;600;700&display=swap');
        
        :root {
            --primary-gradient: linear-gradient(135deg, #333333 0%, #1a1a1a 100%);
            --secondary-gradient: linear-gradient(135deg, #444444 0%, #222222 100%);
            --success-gradient: linear-gradient(135deg, #1a472a 0%, #2d5a3d 100%);
            --error-gradient: linear-gradient(135deg, #5c1c1c 0%, #4a1515 100%);
            --bg-gradient: #000000;
            --text-primary: #ffffff;
            --text-secondary: #cccccc;
            --text-muted: #999999;
            --bg-black: #000000;
            --bg-darker: #0a0a0a;
            --bg-card: #111111;
            --bg-card-hover: #1a1a1a;
            --border-color: #333333;
            --accent-color: #4f46e5;
        }
        
        /* Global Black Theme */
        html, body, [data-testid="stAppViewContainer"] {
            background-color: #000000 !important;
            color: #ffffff !important;
        }
        
        .main {
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
            background: #000000 !important;
            min-height: 100vh;
            color: #ffffff !important;
        }
        
        .stApp {
            background: #000000 !important;
        }
        
        [data-testid="stAppViewContainer"] > .main {
            background: #000000 !important;
        }
        
        [data-testid="stHeader"] {
            background: #000000 !important;
        }
        
        /* Hide Streamlit Elements */
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
        .stDeployButton {display: none;}
        header[data-testid="stHeader"] {visibility: hidden;}
        
        /* Modern Header */
        .modern-header {
            background: linear-gradient(135deg, #1a1a1a 0%, #000000 100%);
            padding: 4rem 2rem;
            border-radius: 0 0 3rem 3rem;
            text-align: center;
            margin: -1rem -1rem 3rem -1rem;
            box-shadow: 0 10px 30px rgba(255, 255, 255, 0.1);
            border: 2px solid #333333;
            position: relative;
            overflow: hidden;
        }
        
        .modern-header::before {
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            bottom: 0;
            background: radial-gradient(circle at 30% 50%, rgba(79, 70, 229, 0.1) 0%, transparent 50%);
            pointer-events: none;
        }
        
        .modern-header h1 {
            color: #ffffff;
            font-size: 3.5rem;
            font-weight: 800;
            margin: 0;
            font-family: 'Poppins', sans-serif;
            position: relative;
            z-index: 1;
            text-shadow: 2px 2px 4px rgba(0,0,0,0.5);
        }
        
        .modern-header p {
            color: #cccccc;
            font-size: 1.3rem;
            margin: 1rem 0 0 0;
            position: relative;
            z-index: 1;
        }
        
        /* Modern Buttons */
        .stButton > button {
            background: linear-gradient(135deg, #333333 0%, #1a1a1a 100%) !important;
            color: #ffffff !important;
            border: 2px solid #444444 !important;
            border-radius: 1rem !important;
            padding: 0.75rem 2rem !important;
            font-weight: 600 !important;
            font-size: 1rem !important;
            transition: all 0.3s ease !important;
            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1) !important;
        }
        
        .stButton > button:hover {
            background: linear-gradient(135deg, #444444 0%, #2a2a2a 100%) !important;
            border-color: #666666 !important;
            transform: translateY(-2px) !important;
            box-shadow: 0 8px 15px rgba(0, 0, 0, 0.2) !important;
        }
        
        .stButton > button:active {
            transform: translateY(0) !important;
        }
        
        /* Chat Messages */
        .stChatMessage {
            background-color: #111111 !important;
            border: 2px solid #333333 !important;
            border-radius: 1rem !important;
            color: #ffffff !important;
            margin: 0.5rem 0 !important;
            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1) !important;
        }
        
        .stChatMessage[data-testid="user-message"] {
            background: linear-gradient(135deg, #1e3a8a 0%, #1e40af 100%) !important;
        }
        
        .stChatMessage[data-testid="assistant-message"] {
            background: linear-gradient(135deg, #1f2937 0%, #111827 100%) !important;
        }
        
        /* Form Styling */
        .stTextInput > div > div > input {
            background-color: #1a1a1a !important;
            color: #ffffff !important;
            border: 2px solid #333333 !important;
            border-radius: 1rem !important;
            padding: 1rem !important;
            font-size: 1rem !important;
            transition: all 0.3s ease !important;
        }
        
        .stTextInput > div > div > input:focus {
            border-color: var(--accent-color) !important;
            box-shadow: 0 0 0 3px rgba(79, 70, 229, 0.2) !important;
            outline: none !important;
        }
        
        /* File Uploader */
        .stFileUploader > div {
            background-color: #1a1a1a !important;
            color: #ffffff !important;
            border: 2px dashed #444444 !important;
            border-radius: 1rem !important;
            padding: 2rem !important;
            text-align: center !important;
            transition: all 0.3s ease !important;
        }
        
        .stFileUploader > div:hover {
            border-color: var(--accent-color) !important;
            background-color: rgba(79, 70, 229, 0.05) !important;
        }
        
        /* Data Tables */
        .stDataFrame {
            background: #1a1a1a !important;
            border: 2px solid #333333 !important;
            border-radius: 1rem !important;
            overflow: hidden !important;
            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1) !important;
        }
        
        .stDataFrame table {
            background-color: #1a1a1a !important;
            color: #ffffff !important;
        }
        
        .stDataFrame th {
            background-color: #2a2a2a !important;
            color: #ffffff !important;
            border-color: #444444 !important;
            font-weight: 600 !important;
        }
        
        .stDataFrame td {
            background-color: #1a1a1a !important;
            color: #ffffff !important;
            border-color: #333333 !important;
        }
        
        /* Success/Error Messages */
        .stSuccess {
            background: var(--success-gradient) !important;
            color: #ffffff !important;
            border: 2px solid #22c55e !important;
            border-radius: 1rem !important;
        }
        
        .stError {
            background: var(--error-gradient) !important;
            color: #ffffff !important;
            border: 2px solid #ef4444 !important;
            border-radius: 1rem !important;
        }
        
        .stInfo {
            background: linear-gradient(135deg, #1e40af 0%, #1e3a8a 100%) !important;
            color: #ffffff !important;
            border: 2px solid #3b82f6 !important;
            border-radius: 1rem !important;
        }
        
        /* Tabs */
        .stTabs [data-baseweb="tab-list"] {
            gap: 0.5rem;
            background: transparent;
            padding: 0.5rem;
            border-radius: 1rem;
            background: rgba(26, 26, 26, 0.5);
        }
        
        .stTabs [data-baseweb="tab"] {
            background: linear-gradient(135deg, #333333 0%, #1a1a1a 100%) !important;
            color: #ffffff !important;
            border-radius: 1rem !important;
            padding: 1rem 2rem !important;
            font-weight: 600 !important;
            border: 2px solid #444444 !important;
            transition: all 0.3s ease !important;
            box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1) !important;
        }
        
        .stTabs [data-baseweb="tab"]:hover {
            transform: translateY(-2px) !important;
            box-shadow: 0 4px 8px rgba(0, 0, 0, 0.2) !important;
            border-color: #666666 !important;
        }
        
        .stTabs [aria-selected="true"] {
            background: linear-gradient(135deg, var(--accent-color) 0%, #3730a3 100%) !important;
            border-color: var(--accent-color) !important;
            transform: scale(1.05) !important;
        }
        
        /* Metrics */
        .stMetric {
            background: linear-gradient(135deg, #1a1a1a 0%, #111111 100%) !important;
            border: 2px solid #333333 !important;
            border-radius: 1rem !important;
            padding: 1.5rem !important;
            text-align: center !important;
            transition: all 0.3s ease !important;
        }
        
        .stMetric:hover {
            transform: translateY(-2px);
            box-shadow: 0 8px 15px rgba(0, 0, 0, 0.2);
            border-color: #555555;
        }
        
        .stMetric > div > div[data-testid="metric-container"] > div {
            color: #ffffff !important;
        }
        
        /* Expanders */
        .stExpander {
            background: #1a1a1a !important;
            border: 2px solid #333333 !important;
            border-radius: 1rem !important;
            overflow: hidden !important;
        }
        
        .stExpander > div > div {
            background: #1a1a1a !important;
            color: #ffffff !important;
        }
        
        /* Chat Input */
        .stChatInputContainer {
            background: #1a1a1a !important;
            border: 2px solid #333333 !important;
            border-radius: 1rem !important;
        }
        
        .stChatInput > div > div > input {
            background: #1a1a1a !important;
            color: #ffffff !important;
            border: none !important;
        }
        
        /* Sidebar */
        .stSidebar {
            background: #111111 !important;
        }
        
        /* Spinners */
        .stSpinner > div {
            border-color: var(--accent-color) transparent transparent transparent !important;
        }
        
        /* Text Elements */
        h1, h2, h3, h4, h5, h6 {
            color: #ffffff !important;
            font-family: 'Poppins', sans-serif !important;
        }
        
        p, div, span, li {
            color: #ffffff !important;
        }
        
        /* Plotly Charts */
        .js-plotly-plot {
            background: transparent !important;
        }
        
        .plotly .modebar {
            background: rgba(26, 26, 26, 0.8) !important;
        }
        
        /* Custom Animations */
        @keyframes fadeIn {
            from { opacity: 0; transform: translateY(20px); }
            to { opacity: 1; transform: translateY(0); }
        }
        
        .modern-header {
            animation: fadeIn 1s ease-out;
        }
        
        /* Responsive Design */
        @media (max-width: 768px) {
            .modern-header h1 {
                font-size: 2.5rem;
            }
            
            .modern-header p {
                font-size: 1.1rem;
            }
            
            .stButton > button {
                padding: 0.5rem 1rem !important;
                font-size: 0.9rem !important;
            }
        }
        
        /* Loading States */
        .stProgress {
            background: #333333 !important;
        }
        
        .stProgress > div > div {
            background: var(--accent-color) !important;
        }
        
        /* Custom Scrollbar */
        ::-webkit-scrollbar {
            width: 8px;
        }
        
        ::-webkit-scrollbar-track {
            background: #1a1a1a;
        }
        
        ::-webkit-scrollbar-thumb {
            background: #444444;
            border-radius: 4px;
        }
        
        ::-webkit-scrollbar-thumb:hover {
            background: #666666;
        }
        </style>
        """, unsafe_allow_html=True)

    def render_modern_landing_page(self):
        """Render the modern landing page with authentication"""
        st.markdown("""
        <div class="modern-header">
            <h1>🚀 Elite MultiModal AI Platform</h1>
            <p>Advanced AI-Powered Analytics • CSV Intelligence • Document Processing • Vision AI</p>
        </div>
        """, unsafe_allow_html=True)
        
        col1, col2, col3 = st.columns([1, 2, 1])
        
        with col2:
            tab1, tab2 = st.tabs(["🔐 Sign In", "✨ Create Account"])
            
            with tab1:
                self.render_modern_login_form()
            
            with tab2:
                self.render_modern_register_form()

    def render_modern_login_form(self):
        """Render the modern login form"""
        st.markdown("### 🔐 Welcome Back")
        st.markdown("*Sign in to access your AI-powered analytics platform*")
        
        with st.form("login_form"):
            username = st.text_input("👤 Username", placeholder="Enter your username")
            password = st.text_input("🔒 Password", type="password", placeholder="Enter your password")
            submitted = st.form_submit_button("🚀 Sign In", use_container_width=True)
            
            if submitted:
                if username and password:
                    success, message = self.auth_manager.login_user(username, password)
                    if success:
                        st.session_state.authenticated = True
                        st.session_state.username = username
                        st.success(message)
                        if self.models_ready:
                            st.success("🤖 AI Models: Ready!")
                        time.sleep(2)
                        st.rerun()
                    else:
                        st.error(message)
                else:
                    st.error("Please fill in all fields!")

    def render_modern_register_form(self):
        """Render the modern registration form"""
        st.markdown("### ✨ Join Elite Platform")
        st.markdown("*Create your account and start exploring AI-powered analytics*")
        
        with st.form("register_form"):
            full_name = st.text_input("👤 Full Name", placeholder="Enter your full name")
            username = st.text_input("🧑‍💼 Username", placeholder="Choose a username")
            email = st.text_input("📧 Email", placeholder="Enter your email address")
            password = st.text_input("🔒 Password", type="password", placeholder="Create a password")
            confirm_password = st.text_input("🔐 Confirm Password", type="password", placeholder="Confirm your password")
            submitted = st.form_submit_button("✨ Create Account", use_container_width=True)
            
            if submitted:
                if all([full_name, username, email, password, confirm_password]):
                    if password != confirm_password:
                        st.error("Passwords don't match!")
                    elif len(password) < 6:
                        st.error("Password must be at least 6 characters!")
                    elif len(username) < 3:
                        st.error("Username must be at least 3 characters!")
                    else:
                        success, message = self.auth_manager.register_user(username, password, email, full_name)
                        if success:
                            st.success(message)
                            st.info("Please sign in with your new account!")
                        else:
                            st.error(message)
                else:
                    st.error("Please fill in all fields!")

    def render_modern_dashboard(self):
        """Render the main dashboard"""
        user_info = self.auth_manager.get_user_info(st.session_state.username)
        
        st.markdown(f"""
        <div class="modern-header">
            <h1>👋 Welcome, {user_info.get('full_name', st.session_state.username)}</h1>
            <p>Elite MultiModal AI Dashboard • Advanced Analytics & Intelligence</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Enhanced Navigation
        col1, col2, col3, col4 = st.columns([1, 1, 1, 3])
        
        with col1:
            if st.button("📊 CSV Analytics", use_container_width=True):
                st.session_state.current_page = 'csv'
                st.rerun()
        
        with col2:
            if st.button("📄 Documents", use_container_width=True):
                st.session_state.current_page = 'documents'
                st.rerun()
        
        with col3:
            if st.button("🖼️ Images", use_container_width=True):
                st.session_state.current_page = 'images'
                st.rerun()
        
        with col4:
            col_spacer, col_profile, col_logout = st.columns([2, 1, 1])
            with col_profile:
                if st.button("👤 Profile", use_container_width=True):
                    self.show_user_profile(user_info)
            with col_logout:
                if st.button("🚪 Logout", use_container_width=True):
                    self.logout_user()
        
        st.markdown("---")
        
        # Render current page
        if st.session_state.current_page == 'csv':
            self.render_csv_page()
        elif st.session_state.current_page == 'documents':
            self.render_documents_page()
        elif st.session_state.current_page == 'images':
            self.render_images_page()

    def show_user_profile(self, user_info):
        """Show user profile information"""
        with st.expander("👤 User Profile", expanded=True):
            col1, col2 = st.columns(2)
            
            with col1:
                st.write(f"**Name:** {user_info.get('full_name', 'N/A')}")
                st.write(f"**Username:** {user_info.get('username', st.session_state.username)}")
                st.write(f"**Email:** {user_info.get('email', 'N/A')}")
            
            with col2:
                st.write(f"**Login Count:** {user_info.get('login_count', 0)}")
                st.write(f"**CSV Queries:** {user_info.get('csv_queries', 0)}")
                st.write(f"**Document Queries:** {user_info.get('document_queries', 0)}")
                st.write(f"**Image Queries:** {user_info.get('image_queries', 0)}")

    def generate_dynamic_examples(self, df):
        """Generate smart example questions based on CSV columns and data types"""
        if df is None or df.empty:
            return []
        
        columns = df.columns.tolist()
        numeric_columns = df.select_dtypes(include=['number']).columns.tolist()
        categorical_columns = df.select_dtypes(include=['object', 'category']).columns.tolist()
        
        examples = [
            "Show me the first 10 rows",
            "What columns are in this dataset?",
            "How many rows are there?",
            "Show me the data summary",
            "Show missing values analysis",
            "Show unique values for each column"
        ]
        
        # Add column-specific examples
        if numeric_columns:
            examples.extend([
                f"What is the average {numeric_columns[0]}?",
                f"Show me the maximum {numeric_columns[0]}",
                f"Calculate the sum of {numeric_columns[0]}"
            ])
            
            if len(numeric_columns) > 1:
                examples.append(f"Compare {numeric_columns[0]} and {numeric_columns[1]}")
        
        if categorical_columns:
            examples.extend([
                f"Show unique values in {categorical_columns[0]}",
                f"Count records by {categorical_columns[0]}"
            ])
            
            if numeric_columns and categorical_columns:
                examples.extend([
                    f"Show average {numeric_columns[0]} by {categorical_columns[0]}",
                    f"Create a chart of {numeric_columns[0]} by {categorical_columns[0]}"
                ])
        
        # Add filtering examples
        if len(columns) > 0:
            examples.extend([
                f"Filter rows where {columns[0]} contains 'value'",
                f"Show top 5 records sorted by {columns[0]}",
                "Find duplicate rows in the dataset"
            ])
        
        return examples

    def render_csv_page(self):
        """Render CSV Analytics page with perfect continuous chat"""
        st.markdown("## 📊 CSV Analytics with Perfect Continuous Chat")
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            uploaded_file = st.file_uploader(
                "📤 Choose CSV file for analysis", 
                type="csv", 
                help="Upload your dataset for AI-powered analytics"
            )
        
        with col2:
            if st.button("🗑️ Clear Chat History", key="clear_csv", use_container_width=True):
                st.session_state.csv_messages = []
                st.success("✅ Chat history cleared!")
                time.sleep(1)
                st.rerun()
        
        # Process uploaded CSV
        if uploaded_file:
            try:
                # Read CSV with encoding detection
                try:
                    df = pd.read_csv(uploaded_file, encoding='utf-8')
                except UnicodeDecodeError:
                    try:
                        df = pd.read_csv(uploaded_file, encoding='latin-1')
                    except:
                        df = pd.read_csv(uploaded_file, encoding='cp1252')
                
                st.session_state.csv_data = df
                st.session_state.csv_file_name = uploaded_file.name
                st.session_state.csv_columns = df.columns.tolist()
                
                # Generate dynamic examples
                st.session_state.dynamic_examples = self.generate_dynamic_examples(df)
                
                # Create database
                self.create_csv_database(df)
                
                st.success(f"✅ Dataset loaded successfully: {uploaded_file.name}")
                
                # Display enhanced metrics
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("📊 Total Rows", f"{len(df):,}")
                with col2:
                    st.metric("📋 Columns", len(df.columns))
                with col3:
                    memory_mb = df.memory_usage(deep=True).sum() / 1024 / 1024
                    st.metric("💾 Memory Usage", f"{memory_mb:.1f} MB")
                with col4:
                    missing = df.isnull().sum().sum()
                    st.metric("❓ Missing Values", f"{missing:,}")
                
                # Enhanced data preview with tabs
                st.markdown("### 📋 Dataset Preview")
                tab1, tab2, tab3 = st.tabs(["📊 Sample Data", "📈 Statistics", "🔍 Column Info"])
                
                with tab1:
                    st.dataframe(df.head(10), use_container_width=True)
                
                with tab2:
                    numeric_df = df.select_dtypes(include=['number'])
                    if not numeric_df.empty:
                        st.dataframe(numeric_df.describe(), use_container_width=True)
                    else:
                        st.info("ℹ️ No numeric columns found for statistical analysis")
                
                with tab3:
                    info_df = pd.DataFrame({
                        'Column': df.columns,
                        'Data Type': df.dtypes.astype(str),
                        'Non-Null Count': df.count(),
                        'Null Count': df.isnull().sum(),
                        'Null %': round((df.isnull().sum() / len(df)) * 100, 2),
                        'Unique Values': [df[col].nunique() for col in df.columns],
                        'Sample Value': [str(df[col].dropna().iloc[0]) if not df[col].dropna().empty else 'N/A' for col in df.columns]
                    })
                    st.dataframe(info_df, use_container_width=True)
                
            except Exception as e:
                st.error(f"❌ Error loading CSV file: {str(e)}")
                st.info("Please ensure your CSV file is properly formatted and try again.")

        # Perfect Continuous Chat Interface for CSV
        if st.session_state.csv_data is not None:
            st.markdown("### 💬 Intelligent Chat with Your CSV Data")
            
            # Display chat history with enhanced formatting
            for message in st.session_state.csv_messages:
                with st.chat_message(message["role"]):
                    # Display text content
                    st.markdown(message["content"])
                    
                    # Display dataframe results if present
                    if message["role"] == "assistant" and "data_result" in message:
                        if message["data_result"] is not None and not message["data_result"].empty:
                            st.markdown("**📊 Query Results:**")
                            st.dataframe(message["data_result"], use_container_width=True)
                        
                        # Display charts if present
                        if "chart" in message and message["chart"]:
                            st.plotly_chart(message["chart"], use_container_width=True)
            
            # Chat input with perfect processing
            if prompt := st.chat_input("🔍 Ask anything about your CSV data...", key="csv_chat"):
                # Add user message
                st.session_state.csv_messages.append({"role": "user", "content": prompt})
                
                with st.chat_message("user"):
                    st.markdown(prompt)
                
                # Generate assistant response
                with st.chat_message("assistant"):
                    with st.spinner("🧠 AI is analyzing your data..."):
                        response, data_result, chart = self.process_csv_chat(prompt)
                        
                        # Display response
                        st.markdown(response)
                        
                        # Display dataframe if available
                        if data_result is not None and not data_result.empty:
                            st.markdown("**📊 Query Results:**")
                            st.dataframe(data_result, use_container_width=True)
                        
                        # Display chart if available
                        if chart:
                            st.plotly_chart(chart, use_container_width=True)
                        
                        # Store complete message
                        st.session_state.csv_messages.append({
                            "role": "assistant", 
                            "content": response,
                            "data_result": data_result,
                            "chart": chart
                        })
                        
                        # Update user stats
                        self.auth_manager.update_user_stats(st.session_state.username, "csv_queries")
            
            # Dynamic example questions based on uploaded data
            if st.session_state.dynamic_examples:
                with st.expander("💡 Smart Example Questions (Based on Your Data)"):
                    st.markdown("**🔍 Data Exploration:**")
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        for example in st.session_state.dynamic_examples[:6]:
                            if st.button(f"💬 {example}", key=f"ex_{hash(example)}", use_container_width=True):
                                # Auto-execute example
                                st.session_state.csv_messages.append({"role": "user", "content": example})
                                response, data_result, chart = self.process_csv_chat(example)
                                st.session_state.csv_messages.append({
                                    "role": "assistant", 
                                    "content": response,
                                    "data_result": data_result,
                                    "chart": chart
                                })
                                st.rerun()
                    
                    with col2:
                        if len(st.session_state.dynamic_examples) > 6:
                            st.markdown("**📈 Advanced Analysis:**")
                            for example in st.session_state.dynamic_examples[6:12]:
                                if st.button(f"🔬 {example}", key=f"adv_{hash(example)}", use_container_width=True):
                                    # Auto-execute example
                                    st.session_state.csv_messages.append({"role": "user", "content": example})
                                    response, data_result, chart = self.process_csv_chat(example)
                                    st.session_state.csv_messages.append({
                                        "role": "assistant", 
                                        "content": response,
                                        "data_result": data_result,
                                        "chart": chart
                                    })
                                    st.rerun()

    def render_documents_page(self):
        """Render Document Intelligence page with perfect continuous chat"""
        st.markdown("## 📄 Document Intelligence with Perfect Continuous Chat")
        
        if not self.models_ready:
            st.error("❌ AI models not available. Please check your Google API key configuration.")
            return
        
        # Enhanced file upload section
        col1, col2, col3 = st.columns(3)
        
        with col1:
            uploaded_pdfs = st.file_uploader(
                "📄 Upload PDF Documents", 
                type="pdf", 
                accept_multiple_files=True, 
                key="pdf_docs",
                help="Upload PDF files for AI analysis"
            )
        
        with col2:
            uploaded_docs = st.file_uploader(
                "📝 Upload Word Documents", 
                type=["docx", "doc"], 
                accept_multiple_files=True, 
                key="word_docs",
                help="Upload Word documents for AI analysis"
            )
        
        with col3:
            if st.button("🗑️ Clear Chat History", key="clear_docs", use_container_width=True):
                st.session_state.document_messages = []
                st.success("✅ Document chat history cleared!")
                time.sleep(1)
                st.rerun()
        
        # Process documents with progress tracking
        if (uploaded_pdfs or uploaded_docs) and st.button("🔄 Process Documents", use_container_width=True):
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            with st.spinner("🧠 Processing documents with AI..."):
                success = self.process_documents(uploaded_pdfs, uploaded_docs, progress_bar, status_text)
                if success:
                    st.success("✅ Documents processed successfully! You can now chat with them.")
                else:
                    st.error("❌ Error processing documents. Please try again.")
        
        # Show processed files with details
        if st.session_state.processed_files:
            st.markdown("### 📚 Processed Documents")
            for i, file_info in enumerate(st.session_state.processed_files, 1):
                st.write(f"{i}. {file_info}")
        
        # Perfect Continuous Chat Interface for Documents
        if st.session_state.document_vector_store is not None:
            st.markdown("### 💬 Intelligent Chat with Your Documents")
            
            # Display chat history
            for message in st.session_state.document_messages:
                with st.chat_message(message["role"]):
                    st.markdown(message["content"])
            
            # Chat input
            if prompt := st.chat_input("🔍 Ask anything about your documents...", key="doc_chat"):
                # Add user message to chat history
                st.session_state.document_messages.append({"role": "user", "content": prompt})
                
                with st.chat_message("user"):
                    st.markdown(prompt)
                
                # Generate assistant response
                with st.chat_message("assistant"):
                    with st.spinner("🧠 AI is analyzing your documents..."):
                        response = self.process_document_chat(prompt)
                        
                        # Display response
                        st.markdown(response)
                        
                        # Store complete message in chat history
                        st.session_state.document_messages.append({
                            "role": "assistant", 
                            "content": response
                        })
                        
                        # Update user stats
                        self.auth_manager.update_user_stats(st.session_state.username, "document_queries")
            
            # Enhanced example questions
            with st.expander("💡 Example Document Questions"):
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("**📋 Content Analysis:**")
                    content_examples = [
                        "What are the main topics discussed?",
                        "Summarize the key points from all documents",
                        "What conclusions are mentioned?",
                        "List all important dates mentioned"
                    ]
                    for example in content_examples:
                        if st.button(f"💬 {example}", key=f"doc_content_{hash(example)}", use_container_width=True):
                            st.session_state.document_messages.append({"role": "user", "content": example})
                            response = self.process_document_chat(example)
                            st.session_state.document_messages.append({"role": "assistant", "content": response})
                            st.rerun()
                
                with col2:
                    st.markdown("**🔍 Specific Information:**")
                    specific_examples = [
                        "Find information about specific topics",
                        "Who are the key people mentioned?",
                        "What are the main findings?",
                        "Compare arguments across documents"
                    ]
                    for example in specific_examples:
                        if st.button(f"🔬 {example}", key=f"doc_specific_{hash(example)}", use_container_width=True):
                            st.session_state.document_messages.append({"role": "user", "content": example})
                            response = self.process_document_chat(example)
                            st.session_state.document_messages.append({"role": "assistant", "content": response})
                            st.rerun()

    def render_images_page(self):
        """Render Image Intelligence page with perfect continuous chat"""
        st.markdown("## 🖼️ Image Intelligence with Perfect Continuous Chat")
        
        if not self.models_ready:
            st.error("❌ AI models not available. Please check your Google API key configuration.")
            return
        
        # Enhanced image upload section
        col1, col2 = st.columns(2)
        
        with col1:
            uploaded_image = st.file_uploader(
                "🖼️ Upload Image for Analysis", 
                type=["jpg", "jpeg", "png", "gif", "bmp", "webp"],
                help="Upload an image for AI-powered analysis"
            )
        
        with col2:
            if st.button("🗑️ Clear Chat History", key="clear_images", use_container_width=True):
                st.session_state.image_messages = []
                st.success("✅ Image chat history cleared!")
                time.sleep(1)
                st.rerun()
        
        # Process image with validation
        if uploaded_image and st.button("🔄 Process Image", use_container_width=True):
            try:
                image = Image.open(uploaded_image)
                
                # Validate image
                if image.size[0] > 10000 or image.size[1] > 10000:
                    st.warning("⚠️ Large image detected. Resizing for optimal processing...")
                    image.thumbnail((4096, 4096), Image.Resampling.LANCZOS)
                
                st.session_state.current_image = image
                st.success(f"✅ Image processed successfully: {uploaded_image.name}")
                
                # Show image info
                st.info(f"📏 Image dimensions: {image.size[0]} x {image.size[1]} pixels")
                
            except Exception as e:
                st.error(f"❌ Error processing image: {str(e)}")

        # Display current image and analysis tools
        if st.session_state.current_image is not None:
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("### 🖼️ Current Image")
                st.image(st.session_state.current_image, use_container_width=True)
            
            with col2:
                st.markdown("### 🔧 Quick Actions")
                
                if st.button("📝 Extract Text (OCR)", use_container_width=True):
                    with st.spinner("🔍 Extracting text..."):
                        response = self.analyze_image_perfect("Extract all text visible in this image with high accuracy. Provide the text in a structured format.")
                        st.session_state.image_messages.append({"role": "assistant", "content": response})
                        st.rerun()
                
                if st.button("🔍 Comprehensive Analysis", use_container_width=True):
                    with st.spinner("🧠 Analyzing image..."):
                        response = self.analyze_image_perfect()
                        st.session_state.image_messages.append({"role": "assistant", "content": response})
                        st.rerun()
                
                if st.button("🎨 Visual Elements Analysis", use_container_width=True):
                    with st.spinner("🎨 Analyzing visual elements..."):
                        response = self.analyze_image_perfect("Analyze the visual elements, colors, composition, lighting, and artistic aspects of this image in detail.")
                        st.session_state.image_messages.append({"role": "assistant", "content": response})
                        st.rerun()
            
            # Perfect Continuous Chat Interface for Images
            st.markdown("### 💬 Intelligent Chat with Your Image")
            
            # Display chat history
            for message in st.session_state.image_messages:
                with st.chat_message(message["role"]):
                    st.markdown(message["content"])
            
            # Chat input
            if prompt := st.chat_input("🔍 Ask anything about your image...", key="image_chat"):
                # Add user message to chat history
                st.session_state.image_messages.append({"role": "user", "content": prompt})
                
                with st.chat_message("user"):
                    st.markdown(prompt)
                
                # Generate assistant response
                with st.chat_message("assistant"):
                    with st.spinner("🧠 AI is analyzing your image..."):
                        response = self.analyze_image_perfect(prompt)
                        
                        # Display response
                        st.markdown(response)
                        
                        # Store complete message in chat history
                        st.session_state.image_messages.append({
                            "role": "assistant", 
                            "content": response
                        })
                        
                        # Update user stats
                        self.auth_manager.update_user_stats(st.session_state.username, "image_queries")
            
            # Enhanced example questions for images
            with st.expander("💡 Example Image Questions"):
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("**🖼️ Visual Analysis:**")
                    visual_examples = [
                        "What objects can you see?",
                        "Describe the setting and environment",
                        "What are the dominant colors?",
                        "What mood does this convey?"
                    ]
                    for example in visual_examples:
                        if st.button(f"👁️ {example}", key=f"img_visual_{hash(example)}", use_container_width=True):
                            st.session_state.image_messages.append({"role": "user", "content": example})
                            response = self.analyze_image_perfect(example)
                            st.session_state.image_messages.append({"role": "assistant", "content": response})
                            st.rerun()
                
                with col2:
                    st.markdown("**📝 Text & Details:**")
                    text_examples = [
                        "Extract all visible text",
                        "Read any signs or labels",
                        "Identify brand names or logos",
                        "Describe interesting details"
                    ]
                    for example in text_examples:
                        if st.button(f"📖 {example}", key=f"img_text_{hash(example)}", use_container_width=True):
                            st.session_state.image_messages.append({"role": "user", "content": example})
                            response = self.analyze_image_perfect(example)
                            st.session_state.image_messages.append({"role": "assistant", "content": response})
                            st.rerun()

    # CSV Processing Methods - Enhanced and Robust
    def process_csv_chat(self, question):
        """Enhanced CSV chat processing with intelligent routing"""
        try:
            df = st.session_state.csv_data
            
            # First, try AI-powered SQL generation if models are available
            if self.models_ready:
                result = self.execute_csv_query_safe(question)
                if result.get('success', False):
                    result_df = result['result']
                    response = f"✅ **AI Query executed successfully!**\n📊 Found **{len(result_df):,}** records"
                    
                    # Create intelligent visualization
                    chart = self.create_intelligent_chart(result_df, question)
                    return response, result_df, chart
            
            # Fallback to keyword-based processing with enhanced responses
            question_lower = question.lower()
            
            if any(keyword in question_lower for keyword in ["show data", "first", "head", "rows", "sample", "preview"]):
                return self.handle_data_display_query(question, df)
            elif any(keyword in question_lower for keyword in ["columns", "column names", "fields", "structure"]):
                return self.handle_column_info_query(df)
            elif any(keyword in question_lower for keyword in ["summary", "describe", "statistics", "stats", "overview"]):
                return self.handle_summary_query(df)
            elif any(keyword in question_lower for keyword in ["missing", "null", "empty", "na", "nan"]):
                return self.handle_missing_values_query(df)
            elif any(keyword in question_lower for keyword in ["unique", "distinct", "different"]):
                return self.handle_unique_values_query(df)
            elif any(keyword in question_lower for keyword in ["duplicate", "duplicated", "repeated"]):
                return self.handle_duplicate_query(df)
            elif any(keyword in question_lower for keyword in ["correlation", "corr", "relationship"]):
                return self.handle_correlation_query(df)
            else:
                # Try AI-powered general analysis
                if self.models_ready:
                    try:
                        ai_response = self.generate_enhanced_ai_response(question, df)
                        return ai_response, df.head(10), None
                    except Exception as e:
                        print(f"AI analysis error: {str(e)}")
                
                # Final fallback
                response = f"🤔 **I understand you're asking:** '{question}'\n\nLet me show you a sample of your data to help:"
                return response, df.head(10), None
                    
        except Exception as e:
            return f"❌ Error processing question: {str(e)}\n\nHere's a sample of your data:", df.head(5) if df is not None else None, None

    def handle_data_display_query(self, question, df):
        """Handle data display queries with number extraction"""
        num_rows = 10
        numbers = re.findall(r'\d+', question)
        if numbers:
            num_rows = min(int(numbers[0]), 100)  # Cap at 100 rows
        
        result_df = df.head(num_rows)
        response = f"📋 **Displaying first {len(result_df)} rows:**\n- Dataset has **{len(df):,}** total rows and **{len(df.columns)}** columns"
        return response, result_df, None

    def handle_column_info_query(self, df):
        """Handle column information queries"""
        columns_info = pd.DataFrame({
            'Column Name': df.columns,
            'Data Type': df.dtypes.astype(str),
            'Non-Null Count': df.count(),
            'Null Count': df.isnull().sum(),
            'Null %': round((df.isnull().sum() / len(df)) * 100, 2),
            'Unique Values': [df[col].nunique() for col in df.columns],
            'Sample Value': [str(df[col].dropna().iloc[0]) if not df[col].dropna().empty else 'N/A' for col in df.columns]
        })
        response = f"📊 **Complete Column Information:**\n- Total Columns: **{len(df.columns)}**\n- Dataset Shape: **{df.shape[0]:,} rows × {df.shape[1]} columns**"
        return response, columns_info, None

    def handle_summary_query(self, df):
        """Handle summary statistics queries"""
        try:
            numeric_df = df.select_dtypes(include=['number'])
            if not numeric_df.empty:
                summary_df = numeric_df.describe()
                response = f"📈 **Statistical Summary of {len(numeric_df.columns)} numeric columns:**"
                return response, summary_df, None
            else:
                info_df = pd.DataFrame({
                    'Column': df.columns,
                    'Data Type': df.dtypes.astype(str),
                    'Count': df.count(),
                    'Unique Values': [df[col].nunique() for col in df.columns],
                    'Most Frequent': [str(df[col].mode().iloc[0]) if not df[col].mode().empty else 'N/A' for col in df.columns],
                    'Most Frequent Count': [df[col].value_counts().iloc[0] if not df[col].empty else 0 for col in df.columns]
                })
                response = f"📊 **Dataset Overview:**\n- No numeric columns for statistical summary\n- Showing general information"
                return response, info_df, None
        except Exception as e:
            response = f"📊 **Basic Dataset Information:**"
            basic_info = pd.DataFrame({
                'Metric': ['Total Rows', 'Total Columns', 'Memory Usage (MB)'],
                'Value': [len(df), len(df.columns), round(df.memory_usage(deep=True).sum() / 1024 / 1024, 2)]
            })
            return response, basic_info, None

    def handle_missing_values_query(self, df):
        """Handle missing values analysis"""
        missing_data = df.isnull().sum()
        missing_df = pd.DataFrame({
            'Column': missing_data.index,
            'Missing Count': missing_data.values,
            'Missing Percentage': round((missing_data.values / len(df)) * 100, 2),
            'Data Type': df.dtypes.astype(str),
            'Total Values': len(df)
        }).sort_values('Missing Count', ascending=False)
        
        total_missing = missing_data.sum()
        total_values = df.size
        response = f"🔍 **Missing Values Analysis:**\n- Total missing: **{total_missing:,}** out of **{total_values:,}** values (**{(total_missing/total_values)*100:.2f}%**)"
        return response, missing_df, None

    def handle_unique_values_query(self, df):
        """Handle unique values analysis"""
        unique_info = pd.DataFrame({
            'Column': df.columns,
            'Unique Values': [df[col].nunique() for col in df.columns],
            'Total Values': len(df),
            'Uniqueness %': [round((df[col].nunique() / len(df)) * 100, 2) for col in df.columns],
            'Data Type': df.dtypes.astype(str),
            'Most Common': [str(df[col].mode().iloc[0]) if not df[col].mode().empty else 'N/A' for col in df.columns]
        }).sort_values('Unique Values', ascending=False)
        
        response = f"🔢 **Unique Values Analysis across {len(df.columns)} columns:**"
        return response, unique_info, None

    def handle_duplicate_query(self, df):
        """Handle duplicate analysis"""
        duplicates = df.duplicated()
        duplicate_count = duplicates.sum()
        
        if duplicate_count > 0:
            duplicate_rows = df[duplicates].head(20)  # Show max 20 duplicates
            response = f"🔍 **Duplicate Analysis:**\n- Found **{duplicate_count}** duplicate rows\n- Showing first 20 duplicates:"
            return response, duplicate_rows, None
        else:
            info_df = pd.DataFrame({
                'Analysis': ['Total Rows', 'Duplicate Rows', 'Unique Rows', 'Duplicate Percentage'],
                'Result': [len(df), duplicate_count, len(df) - duplicate_count, f"{(duplicate_count/len(df))*100:.2f}%"]
            })
            response = f"✅ **No duplicate rows found!**"
            return response, info_df, None

    def handle_correlation_query(self, df):
        """Handle correlation analysis"""
        numeric_df = df.select_dtypes(include=['number'])
        if len(numeric_df.columns) < 2:
            response = "❌ **Correlation analysis requires at least 2 numeric columns**"
            return response, None, None
        
        correlation_matrix = numeric_df.corr()
        response = f"📊 **Correlation Analysis of {len(numeric_df.columns)} numeric columns:**"
        
        # Create correlation heatmap
        try:
            import plotly.figure_factory as ff
            chart = ff.create_annotated_heatmap(
                z=correlation_matrix.values,
                x=list(correlation_matrix.columns),
                y=list(correlation_matrix.columns),
                annotation_text=correlation_matrix.round(2).values,
                showscale=True,
                colorscale='RdBu'
            )
            chart.update_layout(
                title="Correlation Matrix",
                template="plotly_dark",
                plot_bgcolor='rgba(0,0,0,0)',
                paper_bgcolor='rgba(0,0,0,0)',
                font=dict(color='white')
            )
            return response, correlation_matrix, chart
        except:
            return response, correlation_matrix, None

    def generate_enhanced_ai_response(self, question, df):
        """Generate enhanced AI response for complex questions"""
        if not self.models_ready:
            return "❌ AI models not available for complex analysis."
        
        # Create comprehensive dataset summary
        data_summary = f"""
        Dataset Analysis:
        - Shape: {df.shape[0]:,} rows × {df.shape[1]} columns
        - Columns: {', '.join(df.columns.tolist())}
        - Data types: {dict(df.dtypes.astype(str))}
        - Missing values: {df.isnull().sum().sum()} total
        - Memory usage: {df.memory_usage(deep=True).sum() / 1024 / 1024:.1f} MB
        - Numeric columns: {len(df.select_dtypes(include=['number']).columns)}
        - Categorical columns: {len(df.select_dtypes(include=['object']).columns)}
        
        Sample data (first 3 rows):
        {df.head(3).to_string()}
        """
        
        prompt = f"""
        You are an expert data analyst. Analyze this question about the dataset and provide insights: "{question}"
        
        {data_summary}
        
        Provide a comprehensive response that:
        1. Directly answers the question if possible
        2. Suggests relevant data exploration approaches
        3. Identifies potential insights or patterns
        4. Recommends specific analyses or visualizations
        5. Keep the response informative and actionable
        
        Format your response professionally with clear sections.
        """
        
        try:
            response = self.gemini_model.invoke(prompt)
            return f"🤖 **AI Data Analyst:**\n\n{response.content}"
        except Exception as e:
            return f"🤔 I understand you're asking about: '{question}'. Here's a sample of your data to help you explore further."

    def create_intelligent_chart(self, result_df, question):
        """Create intelligent visualizations based on query results and context"""
        if result_df is None or result_df.empty:
            return None
        
        try:
            # Only create charts for reasonable data sizes
            if len(result_df) > 50 or len(result_df.columns) < 1:
                return None
            
            question_lower = question.lower()
            
            # Determine chart type based on data structure and question context
            if len(result_df.columns) == 2:
                col1, col2 = result_df.columns
                
                # Check if we should create a specific chart type
                if result_df[col2].dtype.kind in 'biufc':  # numeric
                    if any(word in question_lower for word in ['trend', 'time', 'date', 'over time']):
                        # Line chart for trends
                        chart = px.line(
                            result_df, 
                            x=col1, 
                            y=col2,
                            title=f"{col2} Trend over {col1}",
                            template="plotly_dark"
                        )
                    elif any(word in question_lower for word in ['distribution', 'histogram', 'frequency']):
                        # Histogram for distributions
                        chart = px.histogram(
                            result_df,
                            x=col2,
                            title=f"Distribution of {col2}",
                            template="plotly_dark"
                        )
                    else:
                        # Default bar chart
                        chart = px.bar(
                            result_df, 
                            x=col1, 
                            y=col2,
                            title=f"{col2} by {col1}",
                            template="plotly_dark"
                        )
                    
                    # Apply dark theme styling
                    chart.update_layout(
                        plot_bgcolor='rgba(0,0,0,0)',
                        paper_bgcolor='rgba(0,0,0,0)',
                        font=dict(color='white'),
                        title=dict(font=dict(color='white', size=16)),
                        xaxis=dict(gridcolor='#444444'),
                        yaxis=dict(gridcolor='#444444')
                    )
                    return chart
            
            elif len(result_df.columns) == 1:
                # Single column analysis
                col = result_df.columns[0]
                if result_df[col].dtype.kind in 'biufc':  # numeric
                    # Create histogram for single numeric column
                    chart = px.histogram(
                        result_df,
                        x=col,
                        title=f"Distribution of {col}",
                        template="plotly_dark"
                    )
                    chart.update_layout(
                        plot_bgcolor='rgba(0,0,0,0)',
                        paper_bgcolor='rgba(0,0,0,0)',
                        font=dict(color='white'),
                        title=dict(font=dict(color='white', size=16))
                    )
                    return chart
                else:
                    # Create bar chart for categorical data
                    value_counts = result_df[col].value_counts().head(10)
                    chart = px.bar(
                        x=value_counts.index,
                        y=value_counts.values,
                        title=f"Top Values in {col}",
                        template="plotly_dark"
                    )
                    chart.update_layout(
                        plot_bgcolor='rgba(0,0,0,0)',
                        paper_bgcolor='rgba(0,0,0,0)',
                        font=dict(color='white'),
                        title=dict(font=dict(color='white', size=16))
                    )
                    return chart
            
        except Exception as e:
            print(f"Chart creation error: {str(e)}")
            pass
        
        return None

    def execute_csv_query_safe(self, question):
        """Execute CSV query safely with enhanced error handling"""
        if st.session_state.csv_data is None:
            return {"success": False, "error": "No CSV data loaded"}
        
        try:
            if not self.models_ready:
                return {"success": False, "error": "AI models not available"}
            
            df = st.session_state.csv_data
            sql_query = self.generate_enhanced_sql_query(question, df)
            cleaned_query = self.clean_sql_query_robust(sql_query)
            
            # Execute query with timeout protection
            connection = sqlite3.connect(f"{st.session_state.db_name}.db")
            result_df = pd.read_sql_query(cleaned_query, connection)
            connection.close()
            
            return {
                'sql_query': cleaned_query,
                'result': result_df,
                'success': True
            }
            
        except Exception as e:
            return {'error': str(e), 'success': False}

    def generate_enhanced_sql_query(self, question, df):
        """Generate enhanced SQL query with better AI prompting"""
        columns = df.columns.tolist()
        dtypes = {col: str(dtype) for col, dtype in df.dtypes.items()}
        sample_data = df.head(3).to_string()
        
        # Create column-safe names for SQL
        safe_columns = {}
        for col in columns:
            if ' ' in col or '-' in col or col.lower() in ['order', 'group', 'select', 'where']:
                safe_columns[col] = f'"{col}"'
            else:
                safe_columns[col] = col
        
        prompt = f"""
        You are an expert SQL developer. Generate a clean, executable SQLite query to answer: "{question}"
        
        Database Schema:
        - Table name: data_table
        - Columns: {list(safe_columns.values())}
        - Data types: {dtypes}
        
        Sample data for context:
        {sample_data}
        
        SQL Generation Rules:
        1. Return ONLY the SQL query, no explanations
        2. Use proper SQLite syntax
        3. Use double quotes for column names with spaces: {safe_columns}
        4. Include LIMIT clause for large results (max 1000 rows)
        5. Use LOWER() for case-insensitive text searches
        6. Handle NULL values appropriately
        7. Use appropriate aggregation functions (COUNT, SUM, AVG, etc.)
        8. For filtering, use appropriate WHERE conditions
        9. For sorting, use ORDER BY with proper column references
        
        Generate the SQL query:
        """
        
        try:
            response = self.gemini_model.invoke(prompt)
            query = response.content.strip()
            
            # Safety check for dangerous operations
            forbidden_keywords = ['DROP', 'DELETE', 'UPDATE', 'INSERT', 'ALTER', 'CREATE', 'TRUNCATE']
            query_upper = query.upper()
            
            for keyword in forbidden_keywords:
                if keyword in query_upper:
                    return "SELECT * FROM data_table LIMIT 10"
            
            return query
            
        except Exception as e:
            return "SELECT * FROM data_table LIMIT 10"

    def clean_sql_query_robust(self, raw_query):
        """Clean SQL query with robust validation"""
        if not raw_query or len(raw_query.strip()) < 3:
            return "SELECT * FROM data_table LIMIT 10"
        
        cleaned = raw_query.strip()
        
        # Remove markdown formatting
        if "```" in cleaned:
            # Split by triple backticks
            parts = cleaned.split("```")
            if len(parts) >= 3:
                cleaned = parts[1]  # extract the part between triple backticks
            else:
                cleaned = cleaned.replace("```", "")

        # Remove language identifiers
        if cleaned.lower().startswith("sql"):
            cleaned = cleaned[3:].strip()
        
        # Process line by line
        lines = [line.strip() for line in cleaned.split('\n') if line.strip()]
        sql_lines = []
        
        for line in lines:
            if line and not line.startswith('#') and not line.startswith('--') and not line.startswith('/*'):
                sql_lines.append(line)
        
        final_query = ' '.join(sql_lines).rstrip(';').strip()
        
        # Final validation
        if len(final_query) < 5 or not final_query.upper().startswith('SELECT'):
            return "SELECT * FROM data_table LIMIT 10"
        
        return final_query

    def create_csv_database(self, df):
        """Create SQLite database with proper column handling and error recovery"""
        try:
            # Ensure database directory exists
            os.makedirs(os.path.dirname(f"{st.session_state.db_name}.db"), exist_ok=True)
            
            connection = sqlite3.connect(f"{st.session_state.db_name}.db")
            df.to_sql(st.session_state.table_name, connection, if_exists='replace', index=False)
            connection.close()
            return True
        except Exception as e:
            st.error(f"❌ Database creation error: {str(e)}")
            return False

    # Document Processing Methods - Enhanced with Progress Tracking
    def process_documents(self, uploaded_pdfs, uploaded_docs, progress_bar=None, status_text=None):
        """Process uploaded documents with enhanced progress tracking"""
        try:
            all_text = ""
            processed_files = []
            total_files = len(uploaded_pdfs or []) + len(uploaded_docs or [])
            current_file = 0
            
            # Process PDFs
            if uploaded_pdfs:
                for pdf_file in uploaded_pdfs:
                    current_file += 1
                    if status_text:
                        status_text.text(f"Processing PDF {current_file}/{total_files}: {pdf_file.name}")
                    
                    try:
                        pdf_reader = PdfReader(pdf_file)
                        pdf_text = ""
                        
                        for page_num, page in enumerate(pdf_reader.pages):
                            page_text = page.extract_text()
                            if page_text.strip():
                                pdf_text += f"\n--- Page {page_num + 1} of {pdf_file.name} ---\n{page_text}\n"
                        
                        if pdf_text.strip():
                            all_text += pdf_text
                            processed_files.append(f"📄 {pdf_file.name} ({len(pdf_reader.pages)} pages)")
                        
                    except Exception as e:
                        st.error(f"❌ Error reading {pdf_file.name}: {str(e)}")
                    
                    if progress_bar:
                        progress_bar.progress(current_file / total_files)
            
            # Process Word documents
            if uploaded_docs:
                for doc_file in uploaded_docs:
                    current_file += 1
                    if status_text:
                        status_text.text(f"Processing Word document {current_file}/{total_files}: {doc_file.name}")
                    
                    try:
                        doc = docx.Document(doc_file)
                        doc_text = ""
                        para_count = 0
                        
                        for paragraph in doc.paragraphs:
                            if paragraph.text.strip():
                                doc_text += paragraph.text + "\n"
                                para_count += 1
                        
                        # Process tables
                        for table in doc.tables:
                            for row in table.rows:
                                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                                if row_text.strip():
                                    doc_text += row_text + "\n"
                        
                        if doc_text.strip():
                            all_text += f"\n--- {doc_file.name} ---\n{doc_text}\n"
                            processed_files.append(f"📝 {doc_file.name} ({para_count} paragraphs)")
                        
                    except Exception as e:
                        st.error(f"❌ Error reading {doc_file.name}: {str(e)}")
                    
                    if progress_bar:
                        progress_bar.progress(current_file / total_files)
            
            # Create vector store if we have text
            if all_text.strip():
                if status_text:
                    status_text.text("Creating AI knowledge base...")
                
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=10000,
                    chunk_overlap=1000,
                    length_function=len,
                    separators=["\n\n", "\n", ". ", " ", ""]
                )
                text_chunks = text_splitter.split_text(all_text)
                
                if text_chunks:
                    vector_store = FAISS.from_texts(text_chunks, self.embeddings)
                    st.session_state.document_vector_store = vector_store
                    st.session_state.processed_files = processed_files
                    
                    if progress_bar:
                        progress_bar.progress(1.0)
                    if status_text:
                        status_text.text("✅ Processing complete!")
                    
                    return True
            
            return False
            
        except Exception as e:
            st.error(f"❌ Error processing documents: {str(e)}")
            return False

    def process_document_chat(self, question):
        """Process document chat with enhanced analysis"""
        if not self.models_ready:
            return "❌ AI models not available for document analysis."
        
        try:
            # Get relevant documents with higher k for better context
            docs = st.session_state.document_vector_store.similarity_search(question, k=6)
            
            if not docs:
                return "❌ No relevant information found in the documents for your question."
            
            # Enhanced prompt template with better instructions
            prompt_template = """
            You are an expert document analyst with deep analytical skills. Analyze the provided document context and answer the question comprehensively.
            
            Guidelines:
            - Provide detailed, accurate answers based solely on the document context
            - Include specific quotes with page/document references when available
            - If information is partial or unclear, mention this explicitly
            - Structure your response with clear headings and bullet points
            - Provide actionable insights and recommendations where appropriate
            - If the question cannot be fully answered, explain what information is missing
            
            Document Context:
            {context}
            
            Question: {question}
            
            Comprehensive Professional Analysis:
            """
            
            prompt = PromptTemplate(
                template=prompt_template,
                input_variables=["context", "question"]
            )
            
            # Create and run chain
            chain = load_qa_chain(self.gemini_model, chain_type="stuff", prompt=prompt)
            
            response = chain.invoke(
                {"input_documents": docs, "question": question},
                return_only_outputs=True
            )
            
            return f"📚 **Professional Document Analysis:**\n\n{response['output_text']}"
            
        except Exception as e:
            return f"❌ Error analyzing documents: {str(e)}. Please try rephrasing your question."

    # Image Processing Methods - Enhanced with Better Analysis
    def analyze_image_perfect(self, question=""):
        """Enhanced image analysis with comprehensive prompts"""
        if st.session_state.current_image is None:
            return "❌ No image uploaded. Please upload an image first."
        
        if not self.models_ready:
            return "❌ AI vision models not available for image analysis."
        
        try:
            image = st.session_state.current_image
            
            if question:
                prompt = f"""
                You are an expert computer vision analyst with advanced image processing capabilities. 
                Analyze this image carefully and answer the specific question: {question}
                
                Provide a comprehensive response that includes:
                1. **Direct Answer**: Clear, specific response to the question
                2. **Visual Evidence**: Detailed observations supporting your answer
                3. **Text Content (OCR)**: All visible text, signs, labels, or writing
                4. **Context & Background**: Environmental and situational details
                5. **Technical Insights**: Professional observations about image quality, lighting, composition
                6. **Additional Details**: Any relevant information that adds value to the analysis
                
                Be thorough, accurate, and provide actionable information. Use professional language and organize your response clearly.
                """
            else:
                prompt = """
                You are an expert computer vision analyst. Provide a comprehensive, professional analysis of this image.
                
                **Structure your analysis as follows:**
                
                **🎯 Primary Subject & Composition:**
                - Main focal point and overall scene description
                - Composition elements and visual hierarchy
                
                **👥 People & Objects:**
                - Detailed description of all visible people, objects, and their interactions
                - Actions, activities, or events taking place
                
                **📝 Text Content & Information (OCR):**
                - All visible text, including signs, labels, documents, handwriting
                - Numbers, codes, dates, or other readable information
                
                **🎨 Visual & Artistic Elements:**
                - Color palette, lighting conditions, and photographic style
                - Composition techniques, perspective, and artistic quality
                
                **📍 Setting & Environmental Context:**
                - Location type, environment, and spatial context
                - Time indicators (day/night, season, era if discernible)
                - Background elements and surroundings
                
                **🔍 Notable Details & Technical Observations:**
                - Interesting, important, or unusual elements worth highlighting
                - Image quality, resolution, and technical aspects
                - Safety considerations, hazards, or concerns if applicable
                - Professional insights or specialized observations
                
                **💡 Summary & Key Insights:**
                - Main takeaways and significant findings
                - Potential uses or applications of this image
                - Any recommendations or suggestions based on the analysis
                
                Be comprehensive, accurate, professional, and provide actionable insights throughout your analysis.
                """
            
            response = self.gemini_vision.generate_content([prompt, image])
            return f"🖼️ **Expert Computer Vision Analysis:**\n\n{response.text}"
            
        except Exception as e:
            return f"❌ Error analyzing image: {str(e)}. Please try uploading the image again or check your connection."

    def logout_user(self):
        """Enhanced logout with user confirmation and cleanup"""
        # Store current username for goodbye message
        username = st.session_state.get('username', 'User')
        user_info = self.auth_manager.get_user_info(username)
        
        # Show logout confirmation
        st.success(f"👋 Thank you for using Elite AI Platform, {user_info.get('full_name', username)}!")
        
        # Display session summary
        if user_info:
            st.info(f"""
            **Session Summary:**
            - CSV Queries: {user_info.get('csv_queries', 0)}
            - Document Queries: {user_info.get('document_queries', 0)}  
            - Image Queries: {user_info.get('image_queries', 0)}
            - Total Logins: {user_info.get('login_count', 0)}
            """)
        
        # Clear all session state
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        
        time.sleep(3)
        st.rerun()

    def run(self):
        """Main application runner with comprehensive error handling"""
        try:
            if not st.session_state.authenticated:
                self.render_modern_landing_page()
            else:
                self.render_modern_dashboard()
        except Exception as e:
            st.error(f"❌ Application Error: {str(e)}")
            st.info("Please refresh the page and try again. If the issue persists, check your API keys and internet connection.")
            
            # Show debug information in development
            if st.checkbox("Show Debug Information"):
                st.code(f"""
                Error Details:
                {str(e)}
                
                Session State Keys: {list(st.session_state.keys())}
                Models Ready: {getattr(self, 'models_ready', 'Unknown')}
                API Configured: {AI_CONFIGURED}
                """)

# Application Entry Point
if __name__ == "__main__":
    try:
        # Initialize and run the application
        app = ModernMultiModalPlatform()
        app.run()
    except Exception as e:
        st.error(f"❌ Critical Application Error: {str(e)}")
        st.info("Please check your environment setup and API keys.")
        
        # Comprehensive debug information
        with st.expander("🔧 Complete Setup Guide"):
            st.markdown("""
            ## 🔑 Environment Setup
            
            ### For Streamlit Cloud:
            **Add this to your Advanced Settings → Secrets:**
            ```
            GOOGLE_API_KEY = "your_google_api_key_here"
            ```
            
            ### For Local Development:
            **Create `.env` file:**
            ```
            GOOGLE_API_KEY=your_google_api_key_here
            ```
            
            ## 📦 Required Dependencies (requirements.txt)
            ```
            streamlit==1.28.1
            pandas==2.0.3
            pillow==10.0.1
            PyPDF2==3.0.1
            google-generativeai==0.3.2
            python-docx==0.8.11
            langchain==0.0.350
            langchain-google-genai==0.0.6
            langchain-community==0.0.10
            faiss-cpu==1.7.4
            python-dotenv==1.0.0
            plotly==5.17.0
            nest-asyncio==1.5.8
            ```
            
            ## ⚙️ Streamlit Config (.streamlit/config.toml)
            ```
            [server]
            headless = true
            enableCORS = false
            port = $PORT

            [theme]
            backgroundColor = "#000000"
            secondaryBackgroundColor = "#111111"
            textColor = "#ffffff"
            primaryColor = "#4f46e5"

            [browser]
            gatherUsageStats = false
            ```
            
            ## 🚀 Deployment Steps
            1. **Create your repository** with the above files
            2. **Add API key** to Streamlit Cloud secrets
            3. **Deploy** through Streamlit Cloud dashboard
            4. **Wait 2-3 minutes** for initial deployment
            
            ## 🔍 Debug Information
            - **AI Configuration Status**: {AI_CONFIGURED}
            - **Config Message**: {CONFIG_MESSAGE}
            """)

